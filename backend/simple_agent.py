"""
Enhanced Multimodal Agent cho QuizForce AI Test - Professional Version
Xử lý toàn diện: text, images, DOCX với images embedded, OCR processing
Được thiết kế theo best practices của AI agents và Python conventions.
"""

import json
import re
import os
import io
import time
import docx
from typing import Dict, Optional, Tuple, Any, List, Union
from PIL import Image
import google.generativeai as genai
from dotenv import load_dotenv
import random
import base64
from pathlib import Path
from dataclasses import dataclass
from enum import Enum
import logging
from zipfile import ZipFile

load_dotenv()

class RequestType(Enum):
    """Enum for different types of API requests."""
    TEXT_ONLY = "text"
    MULTIMODAL = "multimodal"
    IMAGE_ONLY = "image"

@dataclass
class APIRequest:
    """Professional data class for API requests."""
    content: Union[str, List[Any]]
    request_type: RequestType
    retry_count: int = 0
    metadata: Dict[str, Any] = None

@dataclass
class ProcessingResult:
    """Professional result container."""
    success: bool
    data: Any = None
    error: str = None
    metadata: Dict[str, Any] = None

class EnhancedMultimodalQuizAgent:
    """
    Enhanced Professional AI Agent cho Việt Nam - Multimodal Processing:
    1. Text processing với Vietnamese optimization
    2. Image processing với OCR tiên tiến
    3. DOCX processing với embedded images
    4. Professional error handling và logging
    5. Batch processing với intelligent quota management
    6. Type-safe operations với proper abstractions
    
    Được thiết kế theo enterprise-level standards cho giáo dục Việt Nam.
    """
    
    def __init__(self, api_key: str = None, enable_logging: bool = True):
        """Initialize Professional Multimodal Agent."""
        self.api_key = api_key or os.getenv("GOOGLE_API_KEY")
        if not self.api_key:
            raise ValueError("❌ Cần có API key để khởi động agent")
        
        # Setup logging
        self._setup_logging(enable_logging)
        
        # Initialize Gemini với professional configuration
        genai.configure(api_key=self.api_key)
        self.model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Agent metadata
        self.agent_info = {
            "name": "QuizMaster AI Pro Enhanced",
            "version": "4.0",
            "specialization": "Multimodal AI Agent cho Giáo dục Việt Nam",
            "capabilities": [
                "Text Processing", "Image OCR", "DOCX with Images", 
                "Vietnamese Optimization", "Batch Processing", "Smart Recovery"
            ]
        }
        
        # Enhanced configuration
        self.config = {
            "requests_count": 0,
            "last_request_time": 0,
            "base_delay": 0.5,
            "max_retries": 3,
            "batch_size": 10,
            "batch_delay": 5,
            "quota_exceeded_delay": 30,
            "max_image_size": 10 * 1024 * 1024,  # 10MB
            "supported_formats": ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'],
            "max_docx_images": 50,  # Max images per DOCX
        }
        
        # Processing state
        self.failed_questions = []
        self.processing_stats = {
            "total_requests": 0,
            "successful_requests": 0,
            "failed_requests": 0,
            "quota_events": 0,
            "images_processed": 0,
            "docx_images_extracted": 0
        }
        
        self.logger.info(f"✨ {self.agent_info['name']} v{self.agent_info['version']} initialized")
        self.logger.info(f"🎯 {self.agent_info['specialization']}")
        
    def _setup_logging(self, enable: bool):
        """Setup professional logging system."""
        self.logger = logging.getLogger(f"QuizAgent_{id(self)}")
        self.logger.setLevel(logging.INFO if enable else logging.WARNING)
        
        if not self.logger.handlers and enable:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            self.logger.addHandler(handler)
    
    def _smart_rate_limit(self) -> None:
        """Intelligent rate limiting với exponential backoff."""
        current_time = time.time()
        time_since_last = current_time - self.config["last_request_time"]
        
        # Calculate dynamic delay based on request count
        requests_count = self.config["requests_count"]
        if requests_count > 50:
            delay = self.config["base_delay"] * 3
        elif requests_count > 30:
            delay = self.config["base_delay"] * 2
        elif requests_count > 15:
            delay = self.config["base_delay"] * 1.5
        else:
            delay = self.config["base_delay"]
        
        # Apply delay if needed
        if time_since_last < delay:
            sleep_time = delay - time_since_last
            if sleep_time > 0:
                self.logger.info(f"⏳ Rate limiting: waiting {sleep_time:.1f}s")
                time.sleep(sleep_time)
        
        self.config["last_request_time"] = time.time()
        self.config["requests_count"] += 1
    
    def _create_api_request(self, content: Union[str, List[Any]], metadata: Dict = None) -> APIRequest:
        """Create properly typed API request."""
        if isinstance(content, str):
            request_type = RequestType.TEXT_ONLY
        elif isinstance(content, list):
            # Check if it contains images
            has_image = any(isinstance(item, Image.Image) for item in content)
            request_type = RequestType.MULTIMODAL if has_image else RequestType.TEXT_ONLY
        else:
            request_type = RequestType.IMAGE_ONLY
            
        return APIRequest(
            content=content,
            request_type=request_type,
            metadata=metadata or {}
        )
    
    def _execute_api_request(self, request: APIRequest) -> ProcessingResult:
        """Execute API request với comprehensive error handling."""
        try:
            self._smart_rate_limit()
            self.processing_stats["total_requests"] += 1
            
            # Execute based on request type
            if request.request_type == RequestType.TEXT_ONLY:
                response = self.model.generate_content(request.content)
            else:
                # Multimodal request
                response = self.model.generate_content(request.content)
            
            # Extract text from response
            response_text = self._extract_response_text(response)
            
            self.processing_stats["successful_requests"] += 1
            self.logger.debug(f"✅ API request successful: {request.request_type.value}")
            
            return ProcessingResult(
                success=True,
                data=response_text,
                metadata={"request_type": request.request_type.value}
            )
            
        except Exception as e:
            self.processing_stats["failed_requests"] += 1
            error_msg = str(e).lower()
            
            self.logger.error(f"❌ API request failed: {e}")
            
            return ProcessingResult(
                success=False,
                error=str(e),
                metadata={
                    "request_type": request.request_type.value,
                    "error_type": self._classify_error(error_msg)
                }
            )
    
    def _extract_response_text(self, response) -> str:
        """Safely extract text from Gemini response."""
        try:
            # Handle different response types
            if hasattr(response, 'text'):
                return response.text
            elif hasattr(response, 'candidates') and response.candidates:
                candidate = response.candidates[0]
                if hasattr(candidate, 'content'):
                    return candidate.content.parts[0].text
            
            # Fallback
            return str(response)
            
        except Exception as e:
            self.logger.warning(f"⚠️ Response extraction issue: {e}")
            return str(response)
    
    def _classify_error(self, error_str: str) -> str:
        """Classify error type for better handling."""
        if any(keyword in error_str for keyword in ["quota", "limit", "429", "resource_exhausted"]):
            return "quota_exceeded"
        elif "safety" in error_str:
            return "safety_filter"
        elif any(keyword in error_str for keyword in ["invalid", "malformed"]):
            return "invalid_request"
        else:
            return "unknown_error"
    
    def _make_api_request_with_enhanced_recovery(self, 
                                               content: Union[str, List[Any]], 
                                               retries: int = 0, 
                                               metadata: Dict = None) -> str:
        """Enhanced API request với professional recovery mechanism."""
        request = self._create_api_request(content, metadata)
        request.retry_count = retries
        
        result = self._execute_api_request(request)
        
        if result.success:
            return result.data
        
        # Handle specific error types
        error_type = result.metadata.get("error_type", "unknown")
        
        if error_type == "quota_exceeded" and retries < self.config["max_retries"]:
            self.processing_stats["quota_events"] += 1
            
            # Enhanced backoff strategy
            if retries == 0:
                wait_time = 2 + random.uniform(0.5, 1.5)
            elif retries == 1:
                wait_time = 5 + random.uniform(0.5, 1.5)
            else:
                wait_time = self.config["quota_exceeded_delay"] + random.uniform(0.5, 2.0)
            
            self.logger.warning(
                f"⏳ Quota limit, waiting {wait_time:.1f}s before retry {retries + 1}/{self.config['max_retries']}"
            )
            time.sleep(wait_time)
            
            return self._make_api_request_with_enhanced_recovery(content, retries + 1, metadata)
        
        elif error_type == "safety_filter":
            raise Exception("❌ Content bị từ chối bởi safety filter. Vui lòng kiểm tra nội dung.")
        
        elif error_type == "invalid_request":
            raise Exception("❌ Request không hợp lệ. Vui lòng kiểm tra dữ liệu đầu vào.")
        
        else:
            if retries < self.config["max_retries"]:
                wait_time = (2 ** retries) + random.uniform(0.5, 1.5)
                self.logger.warning(f"⚠️ Retrying in {wait_time:.1f}s: {result.error}")
                time.sleep(wait_time)
                return self._make_api_request_with_enhanced_recovery(content, retries + 1, metadata)
            else:
                raise Exception(f"Max retries exceeded: {result.error}")
    
    def process_text_answers(self, answer_text: str) -> Dict[int, str]:
        """Process text answers với enhanced Vietnamese support."""
        self.logger.info("🔍 Processing text answers...")
        
        # Try regex first for efficiency
        regex_result = self._parse_answers_with_regex(answer_text)
        if regex_result:
            self.logger.info(f"✅ Regex parsing successful: {len(regex_result)} answers")
            return regex_result
        
        # Fallback to AI processing
        prompt = self._create_answer_parsing_prompt(answer_text)
        
        try:
            response_text = self._make_api_request_with_enhanced_recovery(
                prompt, 
                metadata={"task": "parse_text_answers"}
            )
            result = self._parse_json_response(response_text)
            
            # Validate and normalize
            validated_result = self._validate_answers(result)
            
            self.logger.info(f"✅ AI parsing successful: {len(validated_result)} valid answers")
            return validated_result
            
        except Exception as e:
            self.logger.error(f"❌ Text answer processing failed: {e}")
            return {}
    
    def process_image_answers(self, image_data: bytes) -> Dict[int, str]:
        """Process image answers với advanced OCR."""
        self.logger.info("🖼️ Processing image answers with OCR...")
        
        try:
            # Validate and optimize image
            image = self._prepare_image_for_processing(image_data)
            self.processing_stats["images_processed"] += 1
            
            # Create OCR prompt
            prompt = self._create_ocr_prompt()
            
            # Make multimodal API request
            response_text = self._make_api_request_with_enhanced_recovery(
                [prompt, image],
                metadata={"task": "ocr_image_answers"}
            )
            
            result = self._parse_json_response(response_text)
            validated_result = self._validate_answers(result)
            
            self.logger.info(f"✅ OCR processing successful: {len(validated_result)} answers extracted")
            return validated_result
            
        except Exception as e:
            self.logger.error(f"❌ Image processing failed: {e}")
            return {}
    
    def _prepare_image_for_processing(self, image_data: bytes) -> Image.Image:
        """Prepare image for optimal processing."""
        if len(image_data) > self.config["max_image_size"]:
            raise Exception(f"❌ File ảnh quá lớn. Tối đa {self.config['max_image_size']/1024/1024:.1f}MB")
        
        image = Image.open(io.BytesIO(image_data))
        
        # Optimize if needed
        if image.width > 2048 or image.height > 2048:
            image.thumbnail((2048, 2048), Image.Resampling.LANCZOS)
            self.logger.info("🔧 Image optimized for processing")
        
        # Convert to RGB if needed
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        return image
    
    def extract_questions_from_docx(self, docx_file) -> Dict[int, str]:
        """Enhanced DOCX extraction với embedded image support."""
        self.logger.info("📄 Extracting questions from DOCX with image support...")
        
        try:
            # Read DOCX content
            file_stream = io.BytesIO(docx_file.getvalue())
            document = docx.Document(file_stream)
            
            # Extract text and structure
            all_content = self._extract_docx_content(document)
            
            # Extract embedded images if any
            embedded_images = self._extract_docx_images(docx_file)
            if embedded_images:
                self.logger.info(f"📷 Found {len(embedded_images)} embedded images in DOCX")
                self.processing_stats["docx_images_extracted"] += len(embedded_images)
            
            # Combine text content
            full_text = "\n".join([content[1] for content in all_content])
            cleaned_text = self._clean_vietnamese_text(full_text)
            
            self.logger.info(f"📋 Extracted {len(cleaned_text)} characters from DOCX")
            
            # Try multiple extraction methods
            question_blocks = self._extract_questions_with_multiple_methods(
                all_content, cleaned_text, embedded_images
            )
            
            self.logger.info(f"✅ Total extracted: {len(question_blocks)} questions")
            return question_blocks
            
        except Exception as e:
            self.logger.error(f"❌ DOCX extraction failed: {e}")
            return {}
    
    def _extract_docx_content(self, document) -> List[Tuple[str, str]]:
        """Extract structured content from DOCX."""
        all_content = []
        
        # Extract paragraphs
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                all_content.append(("paragraph", text))
        
        # Extract tables
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join([
                    cell.text.strip() for cell in row.cells 
                    if cell.text.strip()
                ])
                if row_text:
                    all_content.append(("table", row_text))
        
        return all_content
    
    def _extract_docx_images(self, docx_file) -> List[Image.Image]:
        """Extract embedded images from DOCX file."""
        images = []
        
        try:
            # DOCX is actually a ZIP file
            file_stream = io.BytesIO(docx_file.getvalue())
            
            with ZipFile(file_stream, 'r') as zip_ref:
                # Look for image files in media folder
                image_files = [
                    name for name in zip_ref.namelist() 
                    if name.startswith('word/media/') and 
                    any(name.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'])
                ]
                
                for img_path in image_files[:self.config["max_docx_images"]]:
                    try:
                        img_data = zip_ref.read(img_path)
                        image = Image.open(io.BytesIO(img_data))
                        
                        # Optimize embedded images
                        if image.width > 1024 or image.height > 1024:
                            image.thumbnail((1024, 1024), Image.Resampling.LANCZOS)
                        
                        if image.mode != 'RGB':
                            image = image.convert('RGB')
                            
                        images.append(image)
                        
                    except Exception as e:
                        self.logger.warning(f"⚠️ Couldn't process embedded image {img_path}: {e}")
                        continue
        
        except Exception as e:
            self.logger.warning(f"⚠️ Couldn't extract images from DOCX: {e}")
        
        return images
    
    def _extract_questions_with_multiple_methods(self, 
                                               all_content: List[Tuple[str, str]], 
                                               cleaned_text: str,
                                               embedded_images: List[Image.Image]) -> Dict[int, str]:
        """Extract questions using multiple intelligent methods."""
        question_blocks = {}
        
        # Method 1: Structure-based extraction
        blocks_by_structure = self._extract_by_content_structure(all_content)
        if blocks_by_structure:
            question_blocks.update(blocks_by_structure)
            self.logger.info(f"📋 Structure method: {len(blocks_by_structure)} questions")
        
        # Method 2: Pattern-based extraction
        blocks_by_pattern = self._extract_vietnamese_question_blocks(cleaned_text)
        if len(blocks_by_pattern) > len(question_blocks):
            question_blocks = blocks_by_pattern
            self.logger.info(f"📋 Pattern method: {len(blocks_by_pattern)} questions")
        
        # Method 3: AI-assisted extraction if needed
        if len(question_blocks) < 5:
            ai_blocks = self._ai_extract_questions(cleaned_text[:5000])
            if len(ai_blocks) > len(question_blocks):
                question_blocks = ai_blocks
                self.logger.info(f"📋 AI method: {len(ai_blocks)} questions")
        
        # Method 4: Process embedded images if available
        if embedded_images and len(question_blocks) < 10:
            image_blocks = self._extract_questions_from_images(embedded_images)
            if image_blocks:
                # Merge with existing blocks
                for k, v in image_blocks.items():
                    if k not in question_blocks:
                        question_blocks[k] = v
                self.logger.info(f"📷 Image method added: {len(image_blocks)} questions")
        
        return question_blocks
    
    def _extract_questions_from_images(self, images: List[Image.Image]) -> Dict[int, str]:
        """Extract questions from embedded images."""
        question_blocks = {}
        
        for i, image in enumerate(images):
            try:
                self.logger.info(f"🔍 Processing embedded image {i+1}/{len(images)}")
                
                prompt = """
                Bạn là chuyên gia OCR cho hệ thống giáo dục Việt Nam.
                
                NHIỆM VỤ: Trích xuất câu hỏi trắc nghiệm từ hình ảnh.
                
                YÊU CẦU:
                - Tìm patterns: "Câu X", "X.", "Question X"
                - Trích xuất đầy đủ nội dung câu hỏi và 4 lựa chọn A, B, C, D
                - Format JSON: {"1": "Câu hỏi với lựa chọn đầy đủ..."}
                - Bỏ qua text không phải câu hỏi
                
                CHỈ TRẢ VỀ JSON:
                """
                
                response_text = self._make_api_request_with_enhanced_recovery(
                    [prompt, image],
                    metadata={"task": "extract_questions_from_image", "image_index": i}
                )
                
                result = self._parse_json_response(response_text)
                
                # Validate and add to blocks
                for k, v in result.items():
                    if str(k).isdigit() and len(str(v)) > 50:
                        question_blocks[int(k)] = str(v)
                
            except Exception as e:
                self.logger.warning(f"⚠️ Failed to process embedded image {i+1}: {e}")
                continue
        
        return question_blocks
    
    def _create_answer_parsing_prompt(self, answer_text: str) -> str:
        """Create optimized prompt for answer parsing."""
        return f"""
        Bạn là chuyên gia xử lý đáp án trắc nghiệm cho hệ thống giáo dục Việt Nam.
        
        NHIỆM VỤ: Trích xuất CHÍNH XÁC số câu và đáp án từ văn bản.
        
        QUY TẮC CHUYÊN NGHIỆP:
        - Tìm patterns: "số. đáp_án", "số) đáp_án", "số: đáp_án", "Câu số. đáp_án"
        - Format JSON chuẩn: {{"1": "A", "2": "B", "3": "AC"}}
        - Key là string của số câu
        - Value là đáp án (A,B,C,D hoặc kết hợp AC,BD,ABC,ABCD)
        - Chuẩn hóa đáp án kết hợp theo thứ tự ABC
        - Bỏ qua dòng không chứa đáp án
        
        VĂN BẢN PHÂN TÍCH:
        {answer_text}
        
        VÍ DỤ CHUẨN:
        {{"1": "A", "2": "BD", "3": "C", "10": "ABC", "25": "D"}}
        
        CHỈ TRẢ VỀ JSON:
        """
    
    def _create_ocr_prompt(self) -> str:
        """Create optimized OCR prompt."""
        return """
        Bạn là chuyên gia OCR và xử lý hình ảnh cho hệ thống giáo dục Việt Nam.
        
        NHIỆM VỤ: Đọc và trích xuất đáp án trắc nghiệm từ hình ảnh với độ chính xác cao.
        
        YÊU CẦU CHUYÊN NGHIỆP:
        - Phân tích chính xác từng số câu và đáp án tương ứng
        - Nhận diện cả chữ in và chữ viết tay
        - Format JSON chuẩn: {"1": "A", "2": "B", "3": "AC"}
        - Xử lý cả đáp án đơn (A, B, C, D) và đáp án kép (AC, BD, ABC)
        - Bỏ qua nhiễu, watermark, logo
        - Chuẩn hóa đáp án kết hợp theo thứ tự ABC
        
        PATTERNS TÌM KIẾM:
        - "1. A", "Câu 1: B", "1) AC", "Question 1. D"
        - Bỏ qua text không phải đáp án
        - Chỉ thêm vào kết quả khi chắc chắn
        
        CHỈ TRẢ VỀ JSON:
        """
    
    def _parse_answers_with_regex(self, text: str) -> Dict[int, str]:
        """Enhanced regex parsing for Vietnamese answer patterns."""
        answers = {}
        
        # Vietnamese-optimized patterns
        patterns = [
            r'(?:Câu\s*)?(\d+)\s*\.\s*([A-D]+)',  # Câu 1. A
            r'(?:Câu\s*)?(\d+)\s*\)\s*([A-D]+)',  # Câu 1) A  
            r'(?:Câu\s*)?(\d+)\s*:\s*([A-D]+)',   # Câu 1: A
            r'(?:Câu\s*)?(\d+)\s+([A-D]+)\b',     # Câu 1 A
            r'(?:Question\s*)?(\d+)\s*[\.:\)]\s*([A-D]+)',  # Question patterns
            r'^(\d+)\s*[\.:\)]\s*([A-D]+)',       # 1. A (line start)
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            temp_answers = {}
            
            for match in matches:
                try:
                    q_num = int(match[0])
                    answer = match[1].upper().strip()
                    
                    if re.match(r'^[A-D]+$', answer):
                        # Sort combination answers
                        answer = ''.join(sorted(set(answer)))
                        temp_answers[q_num] = answer
                except (ValueError, IndexError):
                    continue
            
            # Keep best result
            if len(temp_answers) > len(answers):
                answers = temp_answers
        
        return answers
    
    def _validate_answers(self, result: Dict) -> Dict[int, str]:
        """Validate and normalize answer format."""
        validated_result = {}
        
        for k, v in result.items():
            if str(k).isdigit():
                q_num = int(k)
                answer = str(v).upper().strip()
                
                # Validate answer format
                if re.match(r'^[A-D]+$', answer):
                    # Normalize combination answers
                    validated_result[q_num] = ''.join(sorted(set(answer)))
        
        return validated_result
    
    def _parse_json_response(self, response_text: str) -> Dict:
        """Enhanced JSON parsing with multiple fallback strategies."""
        if not response_text:
            return {}
        
        try:
            # Clean response
            cleaned = response_text.strip()
            
            # Remove markdown
            if cleaned.startswith("```json"):
                cleaned = cleaned[7:]
            if cleaned.startswith("```"):
                cleaned = cleaned[3:]
            if cleaned.endswith("```"):
                cleaned = cleaned[:-3]
            
            cleaned = cleaned.strip()
            
            return json.loads(cleaned)
            
        except json.JSONDecodeError:
            self.logger.warning("⚠️ JSON parsing failed, trying extraction methods...")
            
            # Try regex extraction
            json_patterns = [
                r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}',  # Nested objects
                r'\{.*?\}',  # Simple objects
            ]
            
            for pattern in json_patterns:
                matches = re.findall(pattern, response_text, re.DOTALL)
                for match in matches:
                    try:
                        return json.loads(match)
                    except:
                        continue
            
            # Fallback to key-value extraction
            return self._extract_key_values(response_text)
    
    def _extract_key_values(self, text: str) -> Dict:
        """Extract key-value pairs as fallback."""
        result = {}
        
        patterns = [
            r'"(\w+)":\s*"([^"]*)"',
            r"'(\w+)':\s*'([^']*)'",
            r'(\w+):\s*"([^"]*)"',
            r'(\w+):\s*([^,}\n]+)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for key, value in matches:
                result[key] = value.strip()
        
        return result
    
    def _clean_vietnamese_text(self, text: str) -> str:
        """Enhanced Vietnamese text cleaning."""
        # Remove HTML/XML tags
        cleaned = re.sub(r'<[^>]+>', '', text)
        
        # Remove special formatting
        cleaned = re.sub(r'\\[^\\]*\\', '', cleaned)
        cleaned = re.sub(r'\[[^\]]*\]', '', cleaned)
        
        # Normalize whitespace
        cleaned = re.sub(r'\n\s*\n+', '\n\n', cleaned)
        cleaned = re.sub(r'[ \t]+', ' ', cleaned)
        cleaned = re.sub(r' *\n *', '\n', cleaned)
        
        # Keep Vietnamese characters
        vietnamese_chars = r'áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđĐ'
        pattern = f'[^\w\s\.,\-\(\)\[\]{{}}:;""\'\'!?{vietnamese_chars}/\\\\]'
        cleaned = re.sub(pattern, '', cleaned)
        
        # Fix encoding issues
        encoding_fixes = {
            'â€™': "'", 'â€œ': '"', 'â€': '"', 
            'â€"': '-', 'Ã¡': 'á', 'Ã ': 'à',
        }
        
        for old, new in encoding_fixes.items():
            cleaned = cleaned.replace(old, new)
        
        return cleaned.strip()
    
    def _extract_by_content_structure(self, content_list: List[Tuple[str, str]]) -> Dict[int, str]:
        """Extract questions based on content structure."""
        questions = {}
        current_question = None
        current_content = []
        
        question_patterns = [
            r'^(?:Câu\s*)?(\d+)\s*[\.:\)]',
            r'^(?:Question\s*)?(\d+)\s*[\.:\)]',
            r'^(?:Bài\s*)?(\d+)\s*[\.:\)]',
            r'^\d+\.\s*',
            r'^Problem\s*(\d+)',
        ]
        
        for content_type, text in content_list:
            text = text.strip()
            if not text:
                continue
            
            # Check for new question
            found_new = False
            for pattern in question_patterns:
                match = re.match(pattern, text, re.IGNORECASE)
                if match:
                    # Save previous question
                    if current_question and current_content:
                        question_text = "\n".join(current_content)
                        if len(question_text) > 30:
                            questions[current_question] = question_text
                    
                    # Start new question
                    try:
                        current_question = int(match.group(1))
                        current_content = [text]
                        found_new = True
                        break
                    except (ValueError, IndexError):
                        continue
            
            # Add to current question
            if not found_new and current_question:
                current_content.append(text)
                
                # Limit length
                if len("\n".join(current_content)) > 2000:
                    combined = "\n".join(current_content)
                    if len(combined) > 1500:
                        questions[current_question] = combined[:1500] + "..."
                        current_question = None
                        current_content = []
        
        # Save last question
        if current_question and current_content:
            question_text = "\n".join(current_content)
            if len(question_text) > 30:
                questions[current_question] = question_text
        
        return questions
    
    def _extract_vietnamese_question_blocks(self, text: str) -> Dict[int, str]:
        """Enhanced Vietnamese question extraction."""
        question_blocks = {}
        
        patterns = [
            r'Câu\s*(\d+)\s*[\.:]',
            r'Question\s*(\d+)\s*[\.:]',
            r'^(\d+)\s*[\.:]',
            r'\n(\d+)\s*[\.:]',
            r'Bài\s*(\d+)\s*[\.:]',
            r'(\d+)\s*\)',
            r'Problem\s*(\d+)',
        ]
        
        best_result = {}
        
        for pattern in patterns:
            parts = re.split(pattern, text, flags=re.IGNORECASE | re.MULTILINE)
            
            if len(parts) > 2:
                temp_blocks = {}
                
                for i in range(1, len(parts), 2):
                    try:
                        q_num = int(parts[i].strip())
                        if i + 1 < len(parts):
                            q_text = parts[i + 1].strip()
                            
                            if self._validate_question_block(q_text):
                                q_text = self._smart_truncate_question(q_text)
                                temp_blocks[q_num] = q_text
                                
                    except ValueError:
                        continue
                
                if temp_blocks and len(temp_blocks) > len(best_result):
                    best_result = temp_blocks
                    if len(best_result) >= 20:
                        break
        
        return best_result
    
    def _validate_question_block(self, text: str) -> bool:
        """Validate question block quality."""
        if len(text) < 30 or len(text) > 3000:
            return False
        
        choice_pattern = r'[A-D][\.:\)]\s*'
        choices_found = len(re.findall(choice_pattern, text))
        
        return choices_found >= 2
    
    def _smart_truncate_question(self, text: str) -> str:
        """Smart truncation preserving structure."""
        if len(text) <= 1500:
            return text
        
        cut_points = []
        
        # Look for choice D
        d_pattern = r'D[\.:\)][^\n]*'
        d_matches = list(re.finditer(d_pattern, text))
        if d_matches:
            last_d = d_matches[-1]
            cut_points.append(last_d.end())
        
        # Look for sentence endings
        sentence_ends = [m.start() for m in re.finditer(r'[\.!?]\s*(?:\n|$)', text[:1500])]
        cut_points.extend(sentence_ends)
        
        if cut_points:
            best_cut = max([cp for cp in cut_points if cp <= 1500])
            return text[:best_cut].strip()
        
        return text[:1500] + "..."
    
    def _ai_extract_questions(self, text: str) -> Dict[int, str]:
        """AI-assisted question extraction."""
        self.logger.info("🤖 Using AI for question extraction...")
        
        try:
            prompt = f"""
            Bạn là chuyên gia phân tích đề thi trắc nghiệm Việt Nam.
            
            NHIỆM VỤ: Tìm và tách các câu hỏi trắc nghiệm từ văn bản.
            
            QUY TẮC:
            - Tìm patterns: "Câu 1.", "1.", "Question 1"
            - Mỗi câu hỏi phải có đủ 4 lựa chọn A, B, C, D
            - Format JSON: {{"1": "Câu hỏi với lựa chọn...", "2": "Câu hỏi 2..."}}
            - Chỉ trả về câu hỏi có cấu trúc rõ ràng
            
            VĂN BẢN:
            {text[:3000]}
            
            CHỈ TRẢ VỀ JSON:
            """
            
            response = self._make_api_request_with_enhanced_recovery(
                prompt, 
                metadata={"task": "ai_extract_questions"}
            )
            result = self._parse_json_response(response)
            
            questions = {}
            for k, v in result.items():
                if str(k).isdigit() and len(str(v)) > 50:
                    questions[int(k)] = str(v)
            
            return questions
            
        except Exception as e:
            self.logger.warning(f"⚠️ AI extraction failed: {e}")
            return {}
    
    def compile_question_with_image_support(self, question_num: int, question_text: str, 
                                          correct_answer: str, images: List[Dict] = None) -> Dict[str, Any]:
        """Compile question with image support."""
        self.logger.info(f"⚙️ Compiling question {question_num} (images: {len(images) if images else 0})")
        
        prompt = f"""
        Bạn là chuyên gia biên soạn đề thi trắc nghiệm hàng đầu tại Việt Nam.
        
        NHIỆM VỤ: Phân tích và chuẩn hóa câu hỏi trắc nghiệm theo tiêu chuẩn giáo dục Việt Nam.
        
        FORMAT OUTPUT:
        {{
            "so_cau": {question_num},
            "cau_hoi": "Nội dung câu hỏi được làm sạch",
            "lua_chon": {{
                "A": "Lựa chọn A",
                "B": "Lựa chọn B", 
                "C": "Lựa chọn C",
                "D": "Lựa chọn D"
            }},
            "dap_an": "{correct_answer}",
            "do_kho": "trung_binh",
            "mon_hoc": "auto_detect",
            "ghi_chu": "Processed by {self.agent_info['name']}"
        }}
        
        YÊU CẦU:
        - Làm sạch và tách riêng câu hỏi và 4 lựa chọn A, B, C, D
        - Đảm bảo tiếng Việt chuẩn
        - Nếu không đủ 4 lựa chọn, tạo lựa chọn hợp lý
        - Phân loại độ khó và môn học
        
        NỘI DUNG:
        {question_text}
        
        CHỈ TRẢ VỀ JSON:
        """
        
        try:
            response_text = self._make_api_request_with_enhanced_recovery(
                prompt, 
                metadata={"task": "compile_question", "question_num": question_num}
            )
            compiled = self._parse_json_response(response_text)
            
            # Validate and enhance
            if isinstance(compiled, dict) and "cau_hoi" in compiled and "lua_chon" in compiled:
                compiled["so_cau"] = question_num
                compiled["dap_an"] = correct_answer
                
                # Add image info
                if images:
                    compiled["images"] = images
                    compiled["has_images"] = True
                    self.logger.info(f"📷 Added {len(images)} images to question {question_num}")
                else:
                    compiled["has_images"] = False
                
                # Add metadata
                compiled["created_by"] = f"{self.agent_info['name']} v{self.agent_info['version']}"
                compiled["created_time"] = time.strftime("%Y-%m-%d %H:%M:%S")
                
                return compiled
            else:
                return self._create_enhanced_fallback_question(question_num, question_text, correct_answer, images)
                
        except Exception as e:
            self.logger.error(f"❌ Question compilation failed for {question_num}: {e}")
            if "QUOTA_EXCEEDED_AFTER_RETRIES" in str(e):
                raise e
            return self._create_enhanced_fallback_question(question_num, question_text, correct_answer, images, error=str(e))
    
    def _create_enhanced_fallback_question(self, question_num: int, question_text: str, 
                                         correct_answer: str, images: List[Dict] = None, 
                                         error: str = None) -> Dict[str, Any]:
        """Create enhanced fallback question."""
        error_msg = f"[AUTO-GENERATED] {error}" if error else "[AUTO-GENERATED] Needs manual review"
        
        # Smart parsing of question content
        lines = question_text.split('\n')
        question_content = []
        choices = {"A": "Option A", "B": "Option B", "C": "Option C", "D": "Option D"}
        
        current_choice = None
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            choice_match = re.match(r'^([A-D])[\.:\)]\s*(.+)', line)
            if choice_match:
                current_choice = choice_match.group(1)
                choices[current_choice] = choice_match.group(2)
            elif current_choice:
                choices[current_choice] += f" {line}"
            else:
                question_content.append(line)
        
        main_question = "\n".join(question_content) if question_content else question_text[:200] + "..."
        
        fallback = {
            "so_cau": question_num,
            "cau_hoi": main_question,
            "lua_chon": choices,
            "dap_an": correct_answer,
            "do_kho": "can_xu_ly",
            "mon_hoc": "unknown",
            "ghi_chu": f"Auto-generated by {self.agent_info['name']} - Needs review",
            "is_fallback": True,
            "fallback_reason": error_msg,
            "created_time": time.strftime("%Y-%m-%d %H:%M:%S"),
            "has_images": bool(images)
        }
        
        if images:
            fallback["images"] = images
            fallback["image_references"] = [img.get("reference", "Unknown") for img in images]
        
        return fallback
    
    def process_complete_quiz_enhanced(self, answer_data, docx_file, answer_type="text") -> Dict[str, Any]:
        """Complete quiz processing with enhanced features."""
        self.logger.info("🚀 Starting enhanced quiz processing...")
        
        # Reset state
        self.config["requests_count"] = 0
        self.failed_questions = []
        
        results = {
            "success": False,
            "parsed_answers": {},
            "question_blocks": {},
            "compiled_questions": [],
            "errors": [],
            "warnings": [],
            "statistics": {},
            "debug_info": {},
            "agent_info": self.agent_info.copy()
        }
        
        try:
            start_time = time.time()
            
            # Step 1: Process answers
            self.logger.info("📝 Step 1/4: Processing answers...")
            if answer_type == "text":
                results["parsed_answers"] = self.process_text_answers(answer_data)
            else:
                results["parsed_answers"] = self.process_image_answers(answer_data)
            
            if not results["parsed_answers"]:
                results["errors"].append("❌ Could not parse answers. Check format: '1. A', '2. B'")
                return results
            
            # Step 2: Extract questions
            self.logger.info("📄 Step 2/4: Extracting questions from DOCX...")
            results["question_blocks"] = self.extract_questions_from_docx(docx_file)
            
            if not results["question_blocks"]:
                results["errors"].append("❌ Could not extract questions. DOCX must have format: 'Question 1.' or '1.'")
                return results
            
            # Step 3: Mapping and validation
            self.logger.info("🔍 Step 3/4: Data compatibility check...")
            
            answer_keys = set(results["parsed_answers"].keys())
            question_keys = set(results["question_blocks"].keys())
            
            results["debug_info"] = {
                "answer_keys": sorted(list(answer_keys)),
                "question_keys": sorted(list(question_keys)),
                "answer_count": len(answer_keys),
                "question_count": len(question_keys),
                "processing_mode": "enhanced_multimodal_v4"
            }
            
            matching_keys = answer_keys & question_keys
            
            if not matching_keys:
                mapping = self._enhanced_question_mapping(answer_keys, question_keys)
                if mapping:
                    new_answers = {}
                    for q_key, a_key in mapping.items():
                        new_answers[q_key] = results["parsed_answers"][a_key]
                    results["parsed_answers"] = new_answers
                    matching_keys = set(mapping.keys())
                    results["debug_info"]["applied_mapping"] = mapping
                    results["warnings"].append(f"🔄 Applied smart mapping for {len(mapping)} questions")
            
            if not matching_keys:
                results["errors"].append("❌ No matching questions between answers and questions.")
                return results
            
            # Step 4: Enhanced compilation
            self.logger.info("⚙️ Step 4/4: Compiling quiz with enhanced processing...")
            
            matching_list = sorted(list(matching_keys))
            total_questions = len(matching_list)
            
            compiled_count = 0
            for i, q_num in enumerate(matching_list):
                question_text = results["question_blocks"][q_num]
                correct_answer = results["parsed_answers"][q_num]
                
                self.logger.info(f"⚙️ Processing question {i+1}/{total_questions} (ID: {q_num})")
                
                try:
                    images = self._detect_question_images(question_text)
                    compiled = self.compile_question_with_image_support(
                        q_num, question_text, correct_answer, images
                    )
                    results["compiled_questions"].append(compiled)
                    compiled_count += 1
                    
                except Exception as e:
                    self.logger.error(f"❌ Failed to compile question {q_num}: {e}")
                    fallback = self._create_enhanced_fallback_question(
                        q_num, question_text, correct_answer, error=str(e)
                    )
                    results["compiled_questions"].append(fallback)
            
            # Statistics
            processing_time = time.time() - start_time
            results["statistics"] = {
                "total_questions": total_questions,
                "successful_compilations": compiled_count,
                "success_rate": f"{(compiled_count/total_questions*100):.1f}%",
                "processing_time": f"{processing_time:.2f}s",
                "api_requests_used": self.config["requests_count"]
            }
            
            results["success"] = True
            
            self.logger.info(f"🎉 Enhanced processing complete!")
            self.logger.info(f"✅ Success: {compiled_count}/{total_questions} questions")
            
        except Exception as e:
            results["errors"].append(f"❌ System error: {str(e)}")
            self.logger.error(f"❌ Critical error: {e}")
        
        return results
    
    def _enhanced_question_mapping(self, answer_keys: set, question_keys: set) -> Dict[int, int]:
        """Enhanced mapping with multiple strategies."""
        mapping = {}
        
        answer_list = sorted(list(answer_keys))
        question_list = sorted(list(question_keys))
        
        # Strategy 1: Perfect alignment
        if len(answer_list) == len(question_list):
            for i, q_num in enumerate(question_list):
                if i < len(answer_list):
                    mapping[q_num] = answer_list[i]
            return mapping
        
        # Strategy 2: Sequential mapping
        for i, q_num in enumerate(question_list):
            if i < len(answer_list):
                mapping[q_num] = answer_list[i]
        
        return mapping
    
    def _detect_question_images(self, question_text: str) -> List[Dict]:
        """Detect image references in question text."""
        images = []
        
        image_patterns = [
            r'(?:Hình|Ảnh|Figure|Image)\s*(\d+)',
            r'(?:Sơ đồ|Biểu đồ|Đồ thị)\s*(\d+)',
            r'(?:Quan sát|Xem|Nhìn)\s+(?:hình|ảnh|sơ đồ)',
            r'(?:Dựa vào|Theo)\s+(?:hình|ảnh|sơ đồ)',
        ]
        
        for pattern in image_patterns:
            matches = re.findall(pattern, question_text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, str) and match.isdigit():
                    images.append({
                        "reference": f"Hình {match}",
                        "type": "reference",
                        "detected_pattern": pattern
                    })
                elif "hình" in question_text.lower() or "ảnh" in question_text.lower():
                    images.append({
                        "reference": "Image referenced",
                        "type": "reference", 
                        "detected_pattern": pattern
                    })
        
        return images
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get comprehensive processing statistics."""
        return {
            "agent_info": self.agent_info,
            "processing_stats": self.processing_stats,
            "config": self.config,
            "success_rate": (
                self.processing_stats["successful_requests"] / 
                max(self.processing_stats["total_requests"], 1) * 100
            )
        }
    
    # Compatibility methods
    def process_complete_quiz(self, answer_data, docx_file, answer_type="text") -> Dict[str, Any]:
        """Backward compatibility wrapper."""
        return self.process_complete_quiz_enhanced(answer_data, docx_file, answer_type)
    
    def compile_question(self, question_num: int, question_text: str, correct_answer: str) -> Dict[str, Any]:
        """Backward compatibility wrapper."""
        return self.compile_question_with_image_support(question_num, question_text, correct_answer)

# Compatibility alias for existing code
SimpleQuizAgent = EnhancedMultimodalQuizAgent